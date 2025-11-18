import os, datetime, traceback, gc, re
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.shared import Inches
from copy import deepcopy
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
# from docx2pdf import convert  # Removido devido a problemas de COM

from config import config as config_map

app = Flask(__name__)

# Configuração baseada na variável de ambiente FLASK_CONFIG
config_name = os.getenv('FLASK_CONFIG', 'production' if os.getenv('DATABASE_URL') else 'default')
app_config = config_map.get(config_name, config_map['default'])
app.config.from_object(app_config)

# Permite sobreescrever o banco via DATABASE_URL (Render/Railway/Supabase)
database_url = os.getenv('DATABASE_URL')
if database_url:
    if database_url.startswith('postgres://'):
        database_url = database_url.replace('postgres://', 'postgresql+psycopg2://', 1)
    
    # Garante que URLs do Supabase tenham sslmode=require
    if 'supabase.co' in database_url and 'sslmode' not in database_url:
        separator = '&' if '?' in database_url else '?'
        database_url = f"{database_url}{separator}sslmode=require"
    
    app.config['SQLALCHEMY_DATABASE_URI'] = database_url

# Normaliza caminhos relativos para absolutos
upload_folder = app.config.get('UPLOAD_FOLDER', 'static/uploads')
if not os.path.isabs(upload_folder):
    upload_folder = os.path.join(app.root_path, upload_folder)
    app.config['UPLOAD_FOLDER'] = upload_folder
os.makedirs(upload_folder, exist_ok=True)

db_uri = app.config.get('SQLALCHEMY_DATABASE_URI', '')
if db_uri.startswith('sqlite:///'):
    sqlite_path = db_uri.replace('sqlite:///', '', 1)
    if not os.path.isabs(sqlite_path):
        sqlite_path = os.path.join(app.root_path, sqlite_path)
        app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{sqlite_path}"
    sqlite_dir = os.path.dirname(sqlite_path)
    if sqlite_dir:
        os.makedirs(sqlite_dir, exist_ok=True)

db = SQLAlchemy(app)

# Filtro para formatação de datas brasileiras
@app.template_filter('data_br')
def data_br(data_str):
    """Converte data no formato YYYY-MM-DD para DD/MM/YYYY"""
    if not data_str:
        return '-'
    try:
        # Se já está no formato brasileiro, retorna como está
        if '/' in data_str and len(data_str.split('/')) == 3:
            return data_str
        
        # Se está no formato ISO (YYYY-MM-DD), converte para brasileiro
        if '-' in data_str and len(data_str.split('-')) == 3:
            partes = data_str.split('-')
            if len(partes[0]) == 4:  # YYYY-MM-DD
                return f"{partes[2]}/{partes[1]}/{partes[0]}"
            elif len(partes[2]) == 4:  # DD-MM-YYYY
                return f"{partes[0]}/{partes[1]}/{partes[2]}"
        
        return data_str
    except:
        return data_str

@app.template_filter('datetime_br')
def datetime_br(datetime_str):
    """Converte datetime para formato brasileiro DD/MM/YYYY HH:MM"""
    if not datetime_str:
        return '-'
    try:
        # Se contém data e hora separadas por espaço
        if ' ' in datetime_str:
            data_part, hora_part = datetime_str.split(' ', 1)
            data_formatada = data_br(data_part)
            return f"{data_formatada} {hora_part}"
        
        return data_br(datetime_str)
    except:
        return datetime_str

# ---------------------------
# MODELOS DO BANCO
# ---------------------------
class Ocorrencia(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    tipo = db.Column(db.String(100))
    data_fato = db.Column(db.String(50))
    data_acionamento = db.Column(db.String(50))
    link = db.Column(db.String(255))
    endereco = db.Column(db.String(255))
    cidade = db.Column(db.String(100))
    complemento = db.Column(db.String(255))
    coordenada = db.Column(db.String(255))
    nome_vitima = db.Column(db.String(100))
    cpf = db.Column(db.String(50))
    sexo = db.Column(db.String(20))
    idade = db.Column(db.String(10))
    filiacao = db.Column(db.String(255))
    naturalidade = db.Column(db.String(100))
    contatos = db.Column(db.String(255))
    enderecos = db.Column(db.String(255))
    vestimentas = db.Column(db.String(255))
    caracteristicas_vitima = db.Column(db.String(255))
    condicoes_neurologicas = db.Column(db.String(255))
    inf_medicas = db.Column(db.String(255))
    experiencia_e_resistencia = db.Column(db.String(255))
    tipo_terreno_agua = db.Column(db.String(255))
    condicoes_do_tipo_terreno = db.Column(db.String(255))
    condicoes_climaticas = db.Column(db.String(255))
    historico_ocorrencia = db.Column(db.Text)
    finalizada = db.Column(db.Boolean, default=False)

    # campos de imagem iniciais (guarda só o caminho)
    img_condic_metereologica = db.Column(db.String(255))
    img_local = db.Column(db.String(255))
    img_upv = db.Column(db.String(255))
    img_satelite_upv = db.Column(db.String(255))
    img_raio_busca = db.Column(db.String(255))
    img_tab_mare = db.Column(db.String(255))
    img_prev_temp_onda = db.Column(db.String(255))
    
    # flags para indicar se as imagens devem ser utilizadas
    usa_img_condic_metereologica = db.Column(db.Boolean, default=True)
    usa_img_local = db.Column(db.Boolean, default=True)
    usa_img_upv = db.Column(db.Boolean, default=True)
    usa_img_satelite_upv = db.Column(db.Boolean, default=True)
    usa_img_raio_busca = db.Column(db.Boolean, default=True)
    usa_img_tab_mare = db.Column(db.Boolean, default=True)
    usa_img_prev_temp_onda = db.Column(db.Boolean, default=True)

class DiaBusca(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ocorrencia_id = db.Column(db.Integer, db.ForeignKey('ocorrencia.id'))
    tipo_busca = db.Column(db.String(20))  # 'terrestre' ou 'aquatica'
    data = db.Column(db.String(50))
    hora_ini = db.Column(db.String(20))
    hora_fim = db.Column(db.String(20))
    numero_ocorrencia = db.Column(db.String(50))
    condicoes = db.Column(db.String(255))
    temp_inicial = db.Column(db.String(20))
    temp_final = db.Column(db.String(20))
    temp_agua = db.Column(db.String(20))
    guarnicao = db.Column(db.String(255))
    recursos = db.Column(db.String(255))
    historico = db.Column(db.Text)
    status_vitima = db.Column(db.String(50))  # Localizada, Não localizada
    img_tab_mare = db.Column(db.String(255))
    img_prev_temp = db.Column(db.String(255))
    img_traj_buscas = db.Column(db.String(255))

class ImagemCustomizada(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ocorrencia_id = db.Column(db.Integer, db.ForeignKey('ocorrencia.id'), nullable=False)
    titulo = db.Column(db.String(255), nullable=False)  # Descrição/título da imagem
    caminho = db.Column(db.String(255), nullable=False)  # Caminho do arquivo
    ordem = db.Column(db.Integer, default=0)  # Ordem de exibição
    usa_imagem = db.Column(db.Boolean, default=True)  # Flag para usar ou não

class RelatorioFinal(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ocorrencia_id = db.Column(db.Integer, db.ForeignKey('ocorrencia.id'))
    status_vitima = db.Column(db.String(255))
    estado_biologico = db.Column(db.String(255))
    coordenada_vitima = db.Column(db.String(255))
    temp_agua = db.Column(db.String(255))
    estado_corpo = db.Column(db.String(255))
    temp_total_buscas = db.Column(db.String(255))
    efetiv_total = db.Column(db.String(255))
    rec_empreg = db.Column(db.String(255))
    img_corpo = db.Column(db.String(255))
    img_local_corpo = db.Column(db.String(255))
    relato = db.Column(db.Text)

with app.app_context():
    db.create_all()


# ---------------------------
# ROTAS
# ---------------------------

# ---------------------------
# UTILITÁRIOS DE DOCX
# ---------------------------

def get_value_or_default(value, default="Não informado"):
    """Retorna o valor fornecido ou o padrão se estiver vazio/None"""
    if value is None or (isinstance(value, str) and not value.strip()):
        return default
    return str(value)

def log_memory_usage(context=""):
    """Loga o uso de memória para diagnóstico"""
    if PSUTIL_AVAILABLE:
        try:
            mem = psutil.virtual_memory()
            print(f"[memoria] {context} - Uso: {mem.percent:.1f}% ({mem.used / 1024 / 1024:.1f} MB / {mem.total / 1024 / 1024:.1f} MB)")
        except:
            pass

INVALID_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*]')

def sanitize_filename(value, fallback="documento"):
    """Remove caracteres inválidos para nome de arquivo e normaliza espaços"""
    if not value:
        value = fallback
    text = INVALID_FILENAME_CHARS.sub('', value).strip()
    # Substitui múltiplos espaços por um único espaço
    text = re.sub(r'\s+', ' ', text)
    return text or fallback

def build_relatorio_filename(ocorrencia):
    """Gera o nome do arquivo DOCX conforme especificação do usuário"""
    base_nome = ocorrencia.nome_vitima or f"Ocorrencia {ocorrencia.id}"
    nome_limpo = sanitize_filename(base_nome, f"Ocorrencia {ocorrencia.id}")
    return f"Relatório Buscas {nome_limpo}".strip()

def compress_image_on_upload(image_path, max_width=1600, max_height=1600, quality=60):
    """Comprime imagem no upload para reduzir tamanho e uso de memória (ultra-agressivo)"""
    if not PIL_AVAILABLE or not image_path or not os.path.exists(image_path):
        return image_path
    
    try:
        # Verifica tamanho do arquivo original
        original_size = os.path.getsize(image_path) / 1024 / 1024  # MB
        if original_size < 0.5:  # Se já é pequena (< 0.5MB), não precisa comprimir
            return image_path
        
        with Image.open(image_path) as img:
            # Converte para RGB se necessário (remove transparência para JPEG)
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Redimensiona se necessário
            if img.width > max_width or img.height > max_height:
                img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            # Determina formato de saída baseado na extensão original
            ext = os.path.splitext(image_path)[1].lower()
            if ext in ('.jpg', '.jpeg'):
                format = 'JPEG'
                save_path = image_path
            elif ext == '.png':
                # PNG mantém transparência, mas vamos converter para JPEG para economizar
                format = 'JPEG'
                save_path = os.path.splitext(image_path)[0] + '.jpg'
            else:
                # Outros formatos -> JPEG
                format = 'JPEG'
                save_path = os.path.splitext(image_path)[0] + '.jpg'
            
            # Salva comprimido
            img.save(save_path, format, optimize=True, quality=quality)
            
            # Remove arquivo original se foi convertido
            if save_path != image_path and os.path.exists(image_path):
                try:
                    os.remove(image_path)
                except:
                    pass
            
            new_size = os.path.getsize(save_path) / 1024 / 1024  # MB
            print(f"[upload] Imagem comprimida: {original_size:.2f} MB -> {new_size:.2f} MB ({save_path})")
            return save_path
    except Exception as e:
        print(f"[upload] Erro ao comprimir {image_path}: {e}")
        traceback.print_exc()
        return image_path

def optimize_image_for_docx(image_path, max_width=800, max_height=800, quality=50):
    """Otimiza imagem para inserção no DOCX (ultra-agressivo para economizar memória)"""
    if not PIL_AVAILABLE or not image_path or not os.path.exists(image_path):
        return image_path
    
    try:
        # Força garbage collection antes de processar
        gc.collect()
        
        with Image.open(image_path) as img:
            # Converte para RGB se necessário
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Sempre redimensiona para garantir tamanho máximo (ultra-agressivo - 800px max)
            if img.width > max_width or img.height > max_height:
                img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            # Salva em arquivo temporário otimizado com compressão ultra-agressiva
            temp_dir = os.path.join(app.root_path, 'static', 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            temp_path = os.path.join(temp_dir, f"opt_{os.path.basename(image_path)}")
            img.save(temp_path, 'JPEG', optimize=True, quality=quality)
            print(f"[imagem] Imagem otimizada: {image_path} -> {temp_path} ({img.width}x{img.height}, qualidade {quality}%)")
            
            # Libera memória imediatamente
            del img
            gc.collect()
            
            return temp_path
    except Exception as e:
        print(f"[imagem] Erro ao otimizar {image_path}: {e}")
        traceback.print_exc()
        gc.collect()
        return image_path

def _iter_block_items(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    # Incluir rodapés e cabeçalhos
    for section in doc.sections:
        # Rodapé
        if hasattr(section, 'footer') and section.footer:
            for para in section.footer.paragraphs:
                yield para
        # Cabeçalho
        if hasattr(section, 'header') and section.header:
            for para in section.header.paragraphs:
                yield para

def remove_placeholders_from_doc(doc, phrases):
    """Remove completamente parágrafos/células que contenham os placeholders informados."""
    if not phrases:
        return
    for para in list(_iter_block_items(doc)):
        text = para.text or ''
        if any(phrase in text for phrase in phrases):
            para.text = ''

def replace_placeholders(doc: Document, mapping: dict, image_keys: list):
    """Substitui placeholders de texto e insere imagens.
    Regras aceitas no template:
    - {{chave}}
    - substituir_chave
    - [chave]
    Para imagens, o parágrafo contendo qualquer uma das marcas acima para a chave
    será limpo e a imagem inserida no local.
    """
    text_keys = {k: v for k, v in mapping.items() if k not in image_keys}

    # Texto
    for para in _iter_block_items(doc):
        for k, v in text_keys.items():
            tokens = [f"{{{{{k}}}}}", f"substituir_{k}", f"[{k}]"]
            for token in tokens:
                if token in para.text:
                    # Se o valor está vazio ou None, usar "Não informado"
                    valor = get_value_or_default(v)
                    para.text = para.text.replace(token, valor)

    # Imagens - processa uma de cada vez para economizar memória
    for para in list(_iter_block_items(doc)):
        for k in image_keys:
            tokens = [f"{{{{{k}}}}}", f"substituir_{k}", f"inserir_imagem_{k}"]
            if any(t in para.text for t in tokens):
                para.text = ""
                path = mapping.get(k)
                if path:
                    try:
                        log_memory_usage(f"Antes de processar imagem {k}")
                        abs_path = os.path.abspath(path)
                        if not os.path.exists(abs_path):
                            print(f"[imagem] Arquivo não encontrado para chave '{k}': {abs_path}")
                            continue
                        # Otimiza imagem antes de inserir
                        optimized_path = optimize_image_for_docx(abs_path)
                        run = para.add_run()
                        run.add_picture(optimized_path, width=Inches(3))
                        print(f"[imagem] Inserida imagem para chave '{k}': {optimized_path}")
                        # Remove arquivo temporário imediatamente após inserir
                        try:
                            if optimized_path != abs_path and os.path.exists(optimized_path):
                                os.remove(optimized_path)
                        except:
                            pass
                        # Libera memória após inserir
                        gc.collect()
                        log_memory_usage(f"Após processar imagem {k}")
                    except Exception as e:
                        print(f"[imagem] Falha ao inserir '{k}' em {path}: {e}")
                        traceback.print_exc()
                        gc.collect()
                # Não logar se path for None - é esperado para campos opcionais

def append_document(target: Document, source: Document):
    for element in list(source.element.body):
        target.element.body.append(deepcopy(element))

def replace_phrase_map(doc: Document, phrase_map: dict, image_phrase_map: dict, image_title_map: dict = None):
    """Substitui frases completas do tipo 'substituir pelo ... fornecido pelo usuario'
    e insere imagens para 'inserir imagem ...'. Funciona em parágrafos e células de tabela.
    
    Args:
        doc: Documento DOCX
        phrase_map: Mapeamento de frases para valores de texto
        image_phrase_map: Mapeamento de frases para caminhos de imagens
        image_title_map: Mapeamento de frases de imagem para placeholders de título
                         Ex: {'inserir imagem condicoes meteorologicas...': 'substituir pelo titulo condicoes meteorologicas'}
    """
    # Primeiro, remove títulos de imagens que não serão usadas
    if image_title_map:
        # Identifica quais imagens serão usadas (apenas com caminhos válidos)
        images_to_use = {key for key, path in image_phrase_map.items() if path}
        
        # Remove parágrafos com títulos de imagens que não serão usadas
        for para in list(_iter_block_items(doc)):
            text = para.text or ''
            # Verifica se o parágrafo contém um título de imagem que não será usado
            for image_phrase, title_phrase in image_title_map.items():
                if image_phrase not in images_to_use and title_phrase in text:
                    # Remove o parágrafo inteiro se contém o título e a imagem não será usada
                    para.text = ''
                    print(f"[titulo] Removido título '{title_phrase}' pois imagem não será usada")
    
    # texto
    for para in _iter_block_items(doc):
        text = para.text
        if not text:
            continue
        changed = False
        for phrase, value in phrase_map.items():
            if phrase and phrase in text:
                # Se o valor está vazio ou None, usar "Não informado"
                valor = get_value_or_default(value)
                text = text.replace(phrase, valor)
                changed = True
        if changed:
            para.text = text
    
    # imagens - processa uma de cada vez para economizar memória
    for para in list(_iter_block_items(doc)):
        t = para.text or ''
        for phrase, path in image_phrase_map.items():
            if phrase and phrase in t:
                para.text = ''
                if path:
                    try:
                        log_memory_usage(f"Antes de processar imagem frase '{phrase[:30]}...'")
                        abs_path = os.path.abspath(path)
                        if not os.path.exists(abs_path):
                            print(f"[imagem] Arquivo não encontrado para frase '{phrase}': {abs_path}")
                            continue
                        # Otimiza imagem antes de inserir
                        optimized_path = optimize_image_for_docx(abs_path)
                        para.add_run().add_picture(optimized_path, width=Inches(3))
                        print(f"[imagem] Inserida imagem para frase '{phrase}': {optimized_path}")
                        # Remove arquivo temporário imediatamente após inserir
                        try:
                            if optimized_path != abs_path and os.path.exists(optimized_path):
                                os.remove(optimized_path)
                        except:
                            pass
                        # Libera memória após inserir
                        gc.collect()
                        log_memory_usage(f"Após processar imagem frase '{phrase[:30]}...'")
                    except Exception as e:
                        print(f"[imagem] Falha ao inserir para frase '{phrase}' em {path}: {e}")
                        traceback.print_exc()
                        gc.collect()
                # Não logar se path for None - é esperado para campos opcionais

@app.route('/')
def index():
    ocorrencias = Ocorrencia.query.all()
    return render_template('index.html', ocorrencias=ocorrencias)

@app.route('/nova', methods=['GET','POST'])
def nova():
    if request.method == 'POST':
        o = Ocorrencia()
        
        # Dados básicos da ocorrência
        o.tipo = request.form.get('tipo', '')
        o.data_fato = request.form.get('data_fato', '')
        o.data_acionamento = request.form.get('data_acionamento', '')
        o.link = request.form.get('link', '')
        o.endereco = request.form.get('endereco', '')
        o.cidade = request.form.get('cidade', '')
        o.complemento = request.form.get('complemento', '')
        o.coordenada = request.form.get('coordenada', '')
        o.historico_ocorrencia = request.form.get('historico_ocorrencia', '')
        
        # Dados da vítima
        o.nome_vitima = request.form.get('nome_vitima', '')
        o.cpf = request.form.get('cpf', '')
        o.sexo = request.form.get('sexo', '')
        o.idade = request.form.get('idade', '')
        o.filiacao = request.form.get('filiacao', '')
        o.naturalidade = request.form.get('naturalidade', '')
        o.contatos = request.form.get('contatos', '')
        o.enderecos = request.form.get('enderecos', '')
        o.vestimentas = request.form.get('vestimentas', '')
        o.caracteristicas_vitima = request.form.get('caracteristicas_vitima', '')
        o.condicoes_neurologicas = request.form.get('condicoes_neurologicas', '')
        o.inf_medicas = request.form.get('inf_medicas', '')
        o.experiencia_e_resistencia = request.form.get('experiencia_e_resistencia', '')
        
        # Dados do ambiente
        o.tipo_terreno_agua = request.form.get('tipo_terreno_agua', '')
        o.condicoes_do_tipo_terreno = request.form.get('condicoes_do_tipo_terreno', '')
        o.condicoes_climaticas = request.form.get('condicoes_climaticas', '')
        
        # Processar imagens e flags de uso
        campos_imagem = [
            ('img_condic_metereologica', 'usa_img_condic_metereologica'),
            ('img_local', 'usa_img_local'),
            ('img_upv', 'usa_img_upv'),
            ('img_satelite_upv', 'usa_img_satelite_upv'),
            ('img_raio_busca', 'usa_img_raio_busca'),
            ('img_tab_mare', 'usa_img_tab_mare'),
            ('img_prev_temp_onda', 'usa_img_prev_temp_onda')
        ]
        
        for campo_img, campo_flag in campos_imagem:
            # Define o flag de uso (padrão True se não especificado)
            use_image = request.form.get(campo_flag, 'on') == 'on'
            setattr(o, campo_flag, use_image)
            
            # Processa o upload da imagem
            if campo_img in request.files:
                f = request.files[campo_img]
                if f and f.filename:
                    filename = secure_filename(f.filename)
                    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    f.save(path)
                    # Comprime a imagem imediatamente após upload (pode alterar o caminho se converter formato)
                    compressed_path = compress_image_on_upload(path)
                    setattr(o, campo_img, compressed_path)
                    print(f"[upload] Salvo '{campo_img}' em {os.path.abspath(path)}")
                    log_memory_usage(f"Após upload {campo_img}")
        
        # Processar imagens customizadas
        db.session.add(o)
        db.session.flush()  # Para obter o ID da ocorrência
        
        # Limite de imagens customizadas
        MAX_IMAGENS_CUSTOMIZADAS = 10
        
        # Processa imagens customizadas (formato: custom_titulo_0, custom_imagem_0, custom_usa_0, etc.)
        custom_index = 0
        imagens_adicionadas = 0
        while True:
            titulo_key = f'custom_titulo_{custom_index}'
            imagem_key = f'custom_imagem_{custom_index}'
            usa_key = f'custom_usa_{custom_index}'
            
            titulo = request.form.get(titulo_key, '').strip()
            usa_imagem = request.form.get(usa_key, 'on') == 'on'
            
            # Se não tem título, não há mais imagens customizadas
            if not titulo:
                break
            
            # Verifica limite de imagens customizadas
            if imagens_adicionadas >= MAX_IMAGENS_CUSTOMIZADAS:
                flash(f'Limite de {MAX_IMAGENS_CUSTOMIZADAS} imagens customizadas atingido. As imagens adicionais foram ignoradas.', 'warning')
                break
            
            # Verifica se há arquivo de imagem
            if imagem_key in request.files:
                f = request.files[imagem_key]
                if f and f.filename:
                    filename = secure_filename(f.filename)
                    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    f.save(path)
                    compressed_path = compress_image_on_upload(path)
                    
                    # Cria registro de imagem customizada
                    img_custom = ImagemCustomizada(
                        ocorrencia_id=o.id,
                        titulo=titulo,
                        caminho=compressed_path,
                        ordem=custom_index,
                        usa_imagem=usa_imagem
                    )
                    db.session.add(img_custom)
                    imagens_adicionadas += 1
                    print(f"[upload] Imagem customizada salva: '{titulo}' em {compressed_path}")
            
            custom_index += 1
        
        db.session.commit()
        return redirect(url_for('index'))
    
    return render_template('form_ocorrencia.html')

def get_proximo_numero_dia(ocorrencia_id):
    """Retorna o próximo número do dia de busca para uma ocorrência"""
    dias_existentes = DiaBusca.query.filter_by(ocorrencia_id=ocorrencia_id).count()
    return dias_existentes + 1

def get_nome_dia(numero):
    """Retorna o nome do dia baseado no número (1º, 2º, ..., 10º)"""
    nomes = {
        1: "primeiro", 2: "segundo", 3: "terceiro", 4: "quarto", 5: "quinto",
        6: "sexto", 7: "sétimo", 8: "oitavo", 9: "nono", 10: "décimo"
    }
    return nomes.get(numero, f"{numero}º")

@app.route('/ocorrencia/<int:id>/novo_dia', methods=['GET','POST'])
def novo_dia(id):
    if request.method == 'POST':
        # Verificar se já existem 10 dias de busca
        dias_existentes = DiaBusca.query.filter_by(ocorrencia_id=id).count()
        if dias_existentes >= 10:
            flash('Limite máximo de 10 dias de busca atingido!', 'error')
            return redirect(url_for('visualizar', id=id))
        
        # Verifica se está vindo o tipo de busca (primeira etapa de seleção)
        if 'tipo_busca' in request.form and 'data' not in request.form:
            # Primeira requisição - redireciona para o formulário apropriado
            tipo_busca = request.form.get('tipo_busca')
            proximo_numero = get_proximo_numero_dia(id)
            nome_dia = get_nome_dia(proximo_numero)
            return render_template('form_dia.html', id=id, tipo_busca=tipo_busca, proximo_numero=proximo_numero, nome_dia=nome_dia)
        
        # Segunda requisição - salvar os dados
        d = DiaBusca()
        d.ocorrencia_id = id
        d.tipo_busca = request.form.get('tipo_busca', 'aquatica')  # default aquatica para compatibilidade
        d.data = request.form.get('data','')
        d.hora_ini = request.form.get('hora_ini','')
        d.hora_fim = request.form.get('hora_fim','')
        d.numero_ocorrencia = request.form.get('numero_ocorrencia','')
        d.condicoes = request.form.get('condicoes','')
        d.temp_inicial = request.form.get('temp_inicial','')
        d.temp_final = request.form.get('temp_final','')
        d.temp_agua = request.form.get('temp_agua','') if d.tipo_busca == 'aquatica' else ''
        d.guarnicao = request.form.get('guarnicao','')
        d.recursos = request.form.get('recursos','')
        d.historico = request.form.get('historico','')
        d.status_vitima = request.form.get('status_vitima','')
        # imagens
        for campo, form_key in [
            ('img_tab_mare','img_tab_mare'),
            ('img_prev_temp','img_prev_temp'),
            ('img_traj_buscas','img_traj_buscas')
        ]:
            f = request.files.get(form_key)
            if f and f.filename:
                filename = secure_filename(f.filename)
                path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                f.save(path)
                # Comprime a imagem imediatamente após upload (pode alterar o caminho se converter formato)
                compressed_path = compress_image_on_upload(path)
                setattr(d, campo, compressed_path)
                print(f"[upload] Salvo dia '{campo}' em {os.path.abspath(path)}")
                log_memory_usage(f"Após upload dia {campo}")
        db.session.add(d)
        db.session.commit()
        return redirect(url_for('index'))
    
    # Verificar se já existem 10 dias de busca
    dias_existentes = DiaBusca.query.filter_by(ocorrencia_id=id).count()
    if dias_existentes >= 10:
        flash('Limite máximo de 10 dias de busca atingido!', 'error')
        return redirect(url_for('visualizar', id=id))
    
    # Mostra a tela de seleção
    proximo_numero = get_proximo_numero_dia(id)
    nome_dia = get_nome_dia(proximo_numero)
    return render_template('selecionar_tipo_dia.html', ocorrencia_id=id, proximo_numero=proximo_numero, nome_dia=nome_dia)

@app.route('/ocorrencia/<int:id>/finalizar', methods=['GET','POST'])
def finalizar(id):
    if request.method == 'POST':
        r = RelatorioFinal()
        r.ocorrencia_id = id
        allowed_statuses = {'Localizada com vida', 'Localizada em óbito', 'Não localizada'}
        status_vitima = request.form.get('status_vitima','').strip()
        if not status_vitima or status_vitima not in allowed_statuses:
            flash('Selecione um status válido da vítima para finalizar a ocorrência.', 'error')
            return redirect(url_for('finalizar', id=id))
        r.status_vitima = status_vitima
        r.coordenada_vitima = request.form.get('coordenada_vitima','')
        r.temp_agua = request.form.get('temp_agua','')
        estado_corpo = request.form.get('estado_corpo','').strip()
        if status_vitima.lower() == 'localizada em óbito':
            if not estado_corpo:
                flash('Informe o estado do corpo para vítimas localizadas em óbito.', 'error')
                return redirect(url_for('finalizar', id=id))
            r.estado_corpo = estado_corpo
        else:
            r.estado_corpo = ''
        r.temp_total_buscas = request.form.get('temp_total_buscas','')
        r.efetiv_total = request.form.get('efetiv_total','')
        r.rec_empreg = request.form.get('rec_empreg','')
        r.relato = request.form.get('relato','')
        # imagens
        campos_imagem_final = ['img_corpo','img_local_corpo'] if status_vitima.lower() == 'localizada em óbito' else []
        for campo in campos_imagem_final:
            f = request.files.get(campo)
            if f and f.filename:
                filename = secure_filename(f.filename)
                path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                f.save(path)
                # Comprime a imagem imediatamente após upload (pode alterar o caminho se converter formato)
                compressed_path = compress_image_on_upload(path)
                setattr(r, campo, compressed_path)
                print(f"[upload] Salvo final '{campo}' em {os.path.abspath(path)}")
                log_memory_usage(f"Após upload final {campo}")
        db.session.add(r)
        # marca ocorrência como finalizada
        oc = Ocorrencia.query.get(id)
        oc.finalizada = True
        db.session.commit()
        return redirect(url_for('index'))
    return render_template('form_final.html', id=id)

@app.route('/gerar/<int:id>')
def gerar(id):
    try:
        log_memory_usage("Início da geração de relatório")
        oc = Ocorrencia.query.get(id)
        if not oc:
            flash('Ocorrência não encontrada!', 'error')
            return redirect(url_for('index'))
        
        dias = DiaBusca.query.filter_by(ocorrencia_id=id).all()
        rf = RelatorioFinal.query.filter_by(ocorrencia_id=id).first()

        # Normaliza caminhos para absolutos baseados em app.root_path
        def get_template_path(filename):
            path = os.path.join(app.root_path, filename)
            if not os.path.exists(path):
                print(f"[gerar] Template não encontrado: {path}")
            return path
        
        def normalize_image_path(image_path):
            """Normaliza caminho de imagem para absoluto"""
            if not image_path:
                return None
            if os.path.isabs(image_path):
                return image_path if os.path.exists(image_path) else None
            # Tenta caminho relativo ao root_path
            abs_path = os.path.join(app.root_path, image_path)
            if os.path.exists(abs_path):
                return abs_path
            # Tenta caminho relativo ao upload_folder
            abs_path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(image_path))
            return abs_path if os.path.exists(abs_path) else None

        # 1) Introdução
        intro_path = get_template_path('modelo_introducao_buscas.docx')
        if os.path.exists(intro_path):
            # Inicia o documento final a partir do modelo de introdução para preservar cabeçalho/rodapé
            base_doc = Document(intro_path)
            # Substituições diretamente no documento base
            mapping = {k: v for k, v in oc.__dict__.items() if k != '_sa_instance_state'}
            # Normaliza caminhos de imagens
            for img_key in ['img_condic_metereologica','img_local','img_upv','img_satelite_upv',
                           'img_raio_busca','img_tab_mare','img_prev_temp_onda']:
                if img_key in mapping:
                    mapping[img_key] = normalize_image_path(mapping[img_key])
            
            image_keys = [
                'img_condic_metereologica','img_local','img_upv','img_satelite_upv',
                'img_raio_busca','img_tab_mare','img_prev_temp_onda'
            ]
            replace_placeholders(base_doc, mapping, image_keys)
            phrase_map = {
                'substituir pelo tipo fornecido pelo usuario': oc.tipo,
                'substituir pelo data/hora do fato fornecido pelo usuario': oc.data_fato,
                'substituir pelo data/hora de acionamento fornecido pelo usuario': oc.data_acionamento,
                'substituir pelo link fornecido pelo usuario': oc.link,
                'substituir pelo endereço do local fornecido pelo usuario': oc.endereco,
                'substituir pelo cidade do local fornecido pelo usuario': oc.cidade,
                'substituir pelo complemento do local fornecido pelo usuario': oc.complemento,
                'substituir pela coordenada do local fornecido pelo usuario': oc.coordenada,
                'substituir pelo nome da vítima fornecido pelo usuario': oc.nome_vitima,
                'substituir pelo cpf fornecido pelo usuario': oc.cpf,
                'substituir pelo sexo fornecido pelo usuario': oc.sexo,
                'substituir pela idade fornecida pelo usuario': oc.idade,
            }
            # Cria mapa de imagens respeitando os flags de uso
            image_phrase_map = {}
            image_mapping = [
                ('inserir imagem condicoes meteorologicas fornecida pelo usuario', oc.img_condic_metereologica, oc.usa_img_condic_metereologica),
                ('inserir imagem local fornecida pelo usuario', oc.img_local, oc.usa_img_local),
                ('inserir imagem upv fornecida pelo usuario', oc.img_upv, oc.usa_img_upv),
                ('inserir imagem satelite upv fornecida pelo usuario', oc.img_satelite_upv, oc.usa_img_satelite_upv),
                ('inserir imagem raio de busca fornecida pelo usuario', oc.img_raio_busca, oc.usa_img_raio_busca),
                ('inserir imagem tábua de maré fornecida pelo usuario', oc.img_tab_mare, oc.usa_img_tab_mare),
                ('inserir imagem previsão de temperatura e ondas fornecida pelo usuario', oc.img_prev_temp_onda, oc.usa_img_prev_temp_onda),
            ]
            
            # Mapeamento de títulos das imagens (frase de imagem -> placeholder de título)
            image_title_map = {
                'inserir imagem condicoes meteorologicas fornecida pelo usuario': 'substituir pelo titulo condicoes meteorologicas',
                'inserir imagem local fornecida pelo usuario': 'substituir pelo titulo imagem do local',
                'inserir imagem upv fornecida pelo usuario': 'substituir pelo titulo imagem upv',
                'inserir imagem satelite upv fornecida pelo usuario': 'substituir pelo titulo imagem satelite upv',
                'inserir imagem raio de busca fornecida pelo usuario': 'substituir pelo titulo imagem raio de busca',
                'inserir imagem tábua de maré fornecida pelo usuario': 'substituir pelo titulo imagem tábua de maré',
                'inserir imagem previsão de temperatura e ondas fornecida pelo usuario': 'substituir pelo titulo imagem previsão de temperatura e ondas',
            }
            
            # Adiciona os títulos ao phrase_map quando as imagens serão usadas
            for phrase, image_path, use_image in image_mapping:
                if use_image and image_path:
                    normalized_path = normalize_image_path(image_path)
                    if normalized_path:
                        image_phrase_map[phrase] = normalized_path
                        # Adiciona o título ao phrase_map se a imagem será usada
                        if phrase in image_title_map:
                            title_phrase = image_title_map[phrase]
                            # Extrai o título do placeholder (remove "substituir pelo titulo")
                            title_text = title_phrase.replace('substituir pelo titulo ', '').title()
                            phrase_map[title_phrase] = title_text
            
            replace_phrase_map(base_doc, phrase_map, image_phrase_map, image_title_map)
            
            # Processa imagens customizadas usando placeholders no modelo
            imagens_customizadas = ImagemCustomizada.query.filter_by(
                ocorrencia_id=id
            ).order_by(ImagemCustomizada.ordem).all()
            
            if imagens_customizadas:
                # Cria mapeamento de placeholders para imagens customizadas
                custom_image_phrase_map = {}
                custom_image_title_map = {}
                
                for idx, img_custom in enumerate(imagens_customizadas):
                    # Placeholders baseados no índice da imagem
                    title_phrase = f'substituir pelo titulo imagem customizada {idx}'
                    image_phrase = f'inserir imagem customizada {idx} fornecida pelo usuario'
                    normalized_path = normalize_image_path(img_custom.caminho) if img_custom.caminho else None
                    # Sempre mapeia o título para conseguir removê-lo quando a imagem não for usada
                    custom_image_title_map[image_phrase] = title_phrase

                    if img_custom.usa_imagem and normalized_path:
                        custom_image_phrase_map[image_phrase] = normalized_path
                        # Adiciona o título ao phrase_map apenas quando será exibido
                        titulo_img = img_custom.titulo.strip() if img_custom.titulo else ''
                        phrase_map[title_phrase] = titulo_img or f"Imagem Customizada {idx + 1}"
                    else:
                        # Marca explicitamente para remover o placeholder da imagem
                        custom_image_phrase_map[image_phrase] = None
                
                # Processa as imagens customizadas usando replace_phrase_map
                if custom_image_phrase_map or custom_image_title_map:
                    replace_phrase_map(base_doc, phrase_map, custom_image_phrase_map, custom_image_title_map)
        else:
            base_doc = Document()

        # 2) Dias de busca
        # Seleciona o template apropriado baseado no tipo de busca
        if dias:
            for i, dia in enumerate(dias, 1):
                log_memory_usage(f"Processando dia {i} de {len(dias)}")
                # Determina qual template usar baseado no tipo de busca
                tipo_busca = getattr(dia, 'tipo_busca', 'aquatica')  # default aquatica para compatibilidade
                if tipo_busca == 'terrestre':
                    dia_tpl_path = get_template_path('modelo_buscas_terrestre_por_dias.docx')
                else:
                    dia_tpl_path = get_template_path('modelo_buscas_aquaticas_por_dias.docx')
                
                if os.path.exists(dia_tpl_path):
                    ddoc = Document(dia_tpl_path)
                    # Anexa o template do dia ao documento final
                    append_document(base_doc, ddoc)
                    # Libera memória do documento temporário
                    del ddoc
                    gc.collect()
                
                # Obter nome do dia
                nome_dia = get_nome_dia(i)
                # Substituições direto no base_doc
                dmapping = {
                    **{k: v for k, v in oc.__dict__.items() if k != '_sa_instance_state'},
                    'dia_indice': i,
                    'dia_nome': nome_dia,
                    'dia_data': dia.data,
                    'dia_tipo_busca': 'Busca Aquática' if tipo_busca == 'aquatica' else 'Busca Terrestre',
                    'dia_hora_inicio': dia.hora_ini,
                    'dia_hora_fim': dia.hora_fim,
                    'dia_numero_ocorrencia': dia.numero_ocorrencia,
                    'dia_condicoes': dia.condicoes,
                    'dia_temp_inicial': dia.temp_inicial,
                    'dia_temp_final': dia.temp_final,
                    'dia_temp_agua': dia.temp_agua,
                    'dia_guarnicao': dia.guarnicao,
                    'dia_recursos': dia.recursos,
                    'dia_historico': dia.historico,
                    'dia_status_vitima': dia.status_vitima,
                    'img_tab_mare_dia': normalize_image_path(dia.img_tab_mare),
                    'img_prev_temp_dia': normalize_image_path(dia.img_prev_temp),
                    'img_traj_buscas_dia': normalize_image_path(dia.img_traj_buscas),
                }
                image_keys = ['img_tab_mare_dia','img_prev_temp_dia','img_traj_buscas_dia']
                replace_placeholders(base_doc, dmapping, image_keys)
                phrase_map = {
                    'substituir pelo nome do dia fornecido pelo usuario': nome_dia,
                    'substituir pela data do dia fornecido pelo usuario': dia.data,
                    'substituir pelo tipo de busca do dia fornecido pelo usuario': 'Busca Aquática' if tipo_busca == 'aquatica' else 'Busca Terrestre',
                    'substituir pelo horario de inicio fornecido pelo usuario': dia.hora_ini,
                    'substituir pelo horario de fim fornecido pelo usuario': dia.hora_fim,
                    'substituir pelo numero da ocorrencia fornecido pelo usuario': dia.numero_ocorrencia,
                    'substituir pelas condicoes fornecidas pelo usuario': dia.condicoes,
                    'substituir pela temperatura inicial fornecida pelo usuario': dia.temp_inicial,
                    'substituir pela temperatura final fornecida pelo usuario': dia.temp_final,
                    'substituir pela temperatura da agua do dia fornecida pelo usuario': dia.temp_agua,
                    'substituir pela guarnicao fornecida pelo usuario': dia.guarnicao,
                    'substituir pelos recursos fornecidos pelo usuario': dia.recursos,
                    'substituir pelo historico do dia fornecido pelo usuario': dia.historico,
                    'substituir pelo status da vitima fornecido pelo usuario': dia.status_vitima,
                }
                image_phrase_map = {}
                # Mapeamento de títulos das imagens dos dias
                image_title_map_dia = {
                    'inserir imagem tábua de maré do dia fornecida pelo usuario': 'substituir pelo titulo imagem tábua de maré do dia',
                    'inserir imagem de previsao do dia fornecida pelo usuario': 'substituir pelo titulo imagem de previsão do dia',
                    'inserir imagem de trajetos das buscas fornecida pelo usuario': 'substituir pelo titulo imagem de trajetos das buscas',
                }
                for phrase, img_path in [
                    ('inserir imagem tábua de maré do dia fornecida pelo usuario', dia.img_tab_mare),
                    ('inserir imagem de previsao do dia fornecida pelo usuario', dia.img_prev_temp),
                    ('inserir imagem de trajetos das buscas fornecida pelo usuario', dia.img_traj_buscas),
                ]:
                    normalized = normalize_image_path(img_path)
                    if normalized:
                        image_phrase_map[phrase] = normalized
                        # Adiciona o título ao phrase_map se a imagem será usada
                        if phrase in image_title_map_dia:
                            title_phrase = image_title_map_dia[phrase]
                            title_text = title_phrase.replace('substituir pelo titulo ', '').title()
                            phrase_map[title_phrase] = title_text
                replace_phrase_map(base_doc, phrase_map, image_phrase_map, image_title_map_dia)
                # Libera memória após processar cada dia
                gc.collect()
                log_memory_usage(f"Após processar dia {i}")

        # 3) Resultado final
        final_tpl_path = get_template_path('modelo_resultado_final_buscas.docx')
        if os.path.exists(final_tpl_path) and rf:
            fdoc = Document(final_tpl_path)
            append_document(base_doc, fdoc)
            fmapping = {
                **{k: v for k, v in oc.__dict__.items() if k != '_sa_instance_state'},
                **{k: v for k, v in rf.__dict__.items() if k not in ['_sa_instance_state','id','ocorrencia_id']},
            }
            # Normaliza caminhos de imagens do relatório final
            for img_key in ['img_corpo','img_local_corpo']:
                if img_key in fmapping:
                    fmapping[img_key] = normalize_image_path(fmapping[img_key])
            
            image_keys = ['img_corpo','img_local_corpo']
            replace_placeholders(base_doc, fmapping, image_keys)
            phrase_map = {
                'substituir pelo status da vitima fornecido pelo usuario': rf.status_vitima,
                'substituir pela coordenada da vitima fornecida pelo usuario': rf.coordenada_vitima,
                'substituir pela temperatura da agua fornecida pelo usuario': rf.temp_agua,
                'substituir pelo estado do corpo fornecido pelo usuario': rf.estado_corpo,
                'substituir pelo tempo total de buscas fornecido pelo usuario': rf.temp_total_buscas,
                'substituir pelo efetivo total fornecido pelo usuario': rf.efetiv_total,
                'substituir pelos recursos empregados fornecidos pelo usuario': rf.rec_empreg,
                'substituir pelo relato final fornecido pelo usuario': rf.relato,
            }
            image_phrase_map = {}
            # Mapeamento de títulos das imagens do resultado final
            image_title_map_final = {
                'inserir imagem do corpo fornecida pelo usuario': 'substituir pelo titulo imagem do corpo',
                'inserir imagem do local do corpo fornecida pelo usuario': 'substituir pelo titulo imagem do local do corpo',
            }
            for phrase, img_path in [
                ('inserir imagem do corpo fornecida pelo usuario', rf.img_corpo),
                ('inserir imagem do local do corpo fornecida pelo usuario', rf.img_local_corpo),
            ]:
                normalized = normalize_image_path(img_path)
                if normalized:
                    image_phrase_map[phrase] = normalized
                    # Adiciona o título ao phrase_map se a imagem será usada
                    if phrase in image_title_map_final:
                        title_phrase = image_title_map_final[phrase]
                        title_text = title_phrase.replace('substituir pelo titulo ', '').title()
                        phrase_map[title_phrase] = title_text
            replace_phrase_map(base_doc, phrase_map, image_phrase_map, image_title_map_final)
            # Remove completamente o campo de estado biológico do template final
            remove_placeholders_from_doc(base_doc, ['substituir pelo estado biologico fornecido pelo usuario'])

        # Salva o arquivo em um diretório temporário
        output_dir = os.path.join(app.root_path, 'static', 'temp')
        os.makedirs(output_dir, exist_ok=True)
        relatorio_nome_base = build_relatorio_filename(oc)
        nome_docx = f'{relatorio_nome_base}.docx'
        output_path = os.path.join(output_dir, nome_docx)
        
        log_memory_usage("Antes de salvar DOCX")
        print(f"[gerar] Salvando relatório em: {output_path}")
        base_doc.save(output_path)
        
        # Limpa arquivos temporários de imagens otimizadas
        try:
            temp_dir = os.path.join(app.root_path, 'static', 'temp')
            for file in os.listdir(temp_dir):
                if file.startswith('opt_') and file.endswith(('.jpg', '.jpeg', '.png')):
                    temp_file = os.path.join(temp_dir, file)
                    try:
                        os.remove(temp_file)
                    except:
                        pass
        except:
            pass
        
        # Força garbage collection para liberar memória
        gc.collect()
        log_memory_usage("Após salvar DOCX e limpeza")
        
        return send_file(output_path, as_attachment=True, download_name=nome_docx)
    
    except Exception as e:
        print(f"[gerar] Erro ao gerar relatório: {e}")
        traceback.print_exc()
        flash(f'Erro ao gerar relatório: {str(e)}', 'error')
        return redirect(url_for('visualizar', id=id))

@app.route('/visualizar/<int:id>')
def visualizar(id):
    ocorrencia = Ocorrencia.query.get_or_404(id)
    dias_busca = DiaBusca.query.filter_by(ocorrencia_id=id).order_by(DiaBusca.data).all()
    relatorio_final = RelatorioFinal.query.filter_by(ocorrencia_id=id).first()
    imagens_customizadas = ImagemCustomizada.query.filter_by(ocorrencia_id=id).order_by(ImagemCustomizada.ordem).all()
    return render_template('visualizar.html', ocorrencia=ocorrencia, dias_busca=dias_busca, relatorio_final=relatorio_final, imagens_customizadas=imagens_customizadas)

@app.route('/editar/<int:id>', methods=['GET', 'POST'])
def editar(id):
    ocorrencia = Ocorrencia.query.get_or_404(id)
    
    if request.method == 'POST':
        # Dados básicos da ocorrência
        ocorrencia.tipo = request.form.get('tipo', '')
        ocorrencia.data_fato = request.form.get('data_fato', '')
        ocorrencia.data_acionamento = request.form.get('data_acionamento', '')
        ocorrencia.link = request.form.get('link', '')
        ocorrencia.endereco = request.form.get('endereco', '')
        ocorrencia.cidade = request.form.get('cidade', '')
        ocorrencia.complemento = request.form.get('complemento', '')
        ocorrencia.coordenada = request.form.get('coordenada', '')
        ocorrencia.historico_ocorrencia = request.form.get('historico_ocorrencia', '')
        
        # Dados da vítima
        ocorrencia.nome_vitima = request.form.get('nome_vitima', '')
        ocorrencia.cpf = request.form.get('cpf', '')
        ocorrencia.sexo = request.form.get('sexo', '')
        ocorrencia.idade = request.form.get('idade', '')
        ocorrencia.filiacao = request.form.get('filiacao', '')
        ocorrencia.naturalidade = request.form.get('naturalidade', '')
        ocorrencia.contatos = request.form.get('contatos', '')
        ocorrencia.enderecos = request.form.get('enderecos', '')
        ocorrencia.vestimentas = request.form.get('vestimentas', '')
        ocorrencia.caracteristicas_vitima = request.form.get('caracteristicas_vitima', '')
        ocorrencia.condicoes_neurologicas = request.form.get('condicoes_neurologicas', '')
        ocorrencia.inf_medicas = request.form.get('inf_medicas', '')
        ocorrencia.experiencia_e_resistencia = request.form.get('experiencia_e_resistencia', '')
        
        # Dados do ambiente
        ocorrencia.tipo_terreno_agua = request.form.get('tipo_terreno_agua', '')
        ocorrencia.condicoes_do_tipo_terreno = request.form.get('condicoes_do_tipo_terreno', '')
        ocorrencia.condicoes_climaticas = request.form.get('condicoes_climaticas', '')
        
        # Processar flags de uso de imagens e upload (apenas se novas imagens foram enviadas)
        campos_imagem = [
            ('img_condic_metereologica', 'usa_img_condic_metereologica'),
            ('img_local', 'usa_img_local'),
            ('img_upv', 'usa_img_upv'),
            ('img_satelite_upv', 'usa_img_satelite_upv'),
            ('img_raio_busca', 'usa_img_raio_busca'),
            ('img_tab_mare', 'usa_img_tab_mare'),
            ('img_prev_temp_onda', 'usa_img_prev_temp_onda')
        ]
        
        for campo_img, campo_flag in campos_imagem:
            # Atualiza o flag de uso
            use_image = request.form.get(campo_flag, 'on') == 'on'
            setattr(ocorrencia, campo_flag, use_image)
            
            # Processa novo upload se houver
            if campo_img in request.files:
                f = request.files[campo_img]
                if f and f.filename:
                    filename = secure_filename(f.filename)
                    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    f.save(path)
                    # Comprime a imagem imediatamente após upload (pode alterar o caminho se converter formato)
                    compressed_path = compress_image_on_upload(path)
                    setattr(ocorrencia, campo_img, compressed_path)
                    print(f"[upload] Atualizado '{campo_img}' em {os.path.abspath(path)}")
                    log_memory_usage(f"Após atualizar {campo_img}")
        
        # Processar imagens customizadas (atualizar existentes e adicionar novas)
        # Primeiro, remove imagens customizadas que foram marcadas para exclusão
        imagens_existentes = ImagemCustomizada.query.filter_by(ocorrencia_id=id).all()
        for img_existente in imagens_existentes:
            delete_key = f'custom_delete_{img_existente.id}'
            if request.form.get(delete_key) == 'on':
                # Remove arquivo físico
                if img_existente.caminho and os.path.exists(img_existente.caminho):
                    try:
                        os.remove(img_existente.caminho)
                    except:
                        pass
                db.session.delete(img_existente)
                continue
            
            # Atualiza título e flag de uso se fornecido
            titulo_key = f'custom_titulo_existente_{img_existente.id}'
            usa_key = f'custom_usa_existente_{img_existente.id}'
            if titulo_key in request.form:
                img_existente.titulo = request.form.get(titulo_key, '').strip()
            if usa_key in request.form:
                img_existente.usa_imagem = request.form.get(usa_key, 'off') == 'on'
            
            # Atualiza imagem se novo arquivo foi enviado
            imagem_key = f'custom_imagem_existente_{img_existente.id}'
            if imagem_key in request.files:
                f = request.files[imagem_key]
                if f and f.filename:
                    # Remove arquivo antigo
                    if img_existente.caminho and os.path.exists(img_existente.caminho):
                        try:
                            os.remove(img_existente.caminho)
                        except:
                            pass
                    # Salva novo arquivo
                    filename = secure_filename(f.filename)
                    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    f.save(path)
                    compressed_path = compress_image_on_upload(path)
                    img_existente.caminho = compressed_path
        
        # Adiciona novas imagens customizadas
        # Limite de imagens customizadas
        MAX_IMAGENS_CUSTOMIZADAS = 10
        
        # Conta quantas imagens customizadas já existem (não deletadas)
        imagens_existentes = ImagemCustomizada.query.filter_by(ocorrencia_id=id).all()
        imagens_nao_deletadas = sum(1 for img in imagens_existentes 
                                   if not request.form.get(f'custom_delete_{img.id}'))
        imagens_restantes = MAX_IMAGENS_CUSTOMIZADAS - imagens_nao_deletadas
        
        custom_index = 0
        imagens_adicionadas = 0
        while True:
            titulo_key = f'custom_titulo_{custom_index}'
            imagem_key = f'custom_imagem_{custom_index}'
            usa_key = f'custom_usa_{custom_index}'
            
            titulo = request.form.get(titulo_key, '').strip()
            usa_imagem = request.form.get(usa_key, 'on') == 'on'
            
            if not titulo:
                break
            
            # Verifica limite de imagens customizadas
            if imagens_adicionadas >= imagens_restantes:
                flash(f'Limite de {MAX_IMAGENS_CUSTOMIZADAS} imagens customizadas atingido. As imagens adicionais foram ignoradas.', 'warning')
                break
            
            if imagem_key in request.files:
                f = request.files[imagem_key]
                if f and f.filename:
                    filename = secure_filename(f.filename)
                    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    f.save(path)
                    compressed_path = compress_image_on_upload(path)
                    
                    # Encontra a maior ordem atual
                    from sqlalchemy import func
                    max_ordem = db.session.query(func.max(ImagemCustomizada.ordem)).filter_by(ocorrencia_id=id).scalar() or -1
                    
                    img_custom = ImagemCustomizada(
                        ocorrencia_id=id,
                        titulo=titulo,
                        caminho=compressed_path,
                        ordem=max_ordem + 1,
                        usa_imagem=usa_imagem
                    )
                    db.session.add(img_custom)
                    imagens_adicionadas += 1
            
            custom_index += 1
        
        db.session.commit()
        flash('Ocorrência atualizada com sucesso!', 'success')
        return redirect(url_for('visualizar', id=id))
    
    imagens_customizadas = ImagemCustomizada.query.filter_by(ocorrencia_id=id).order_by(ImagemCustomizada.ordem).all()
    return render_template('editar.html', ocorrencia=ocorrencia, imagens_customizadas=imagens_customizadas)

@app.route('/editar_dia/<int:id>', methods=['GET', 'POST'])
def editar_dia(id):
    dia = DiaBusca.query.get_or_404(id)
    ocorrencia = Ocorrencia.query.get_or_404(dia.ocorrencia_id)
    
    # Obter o número do dia baseado na ordem de criação
    dias_anteriores = DiaBusca.query.filter(
        DiaBusca.ocorrencia_id == dia.ocorrencia_id,
        DiaBusca.id < dia.id
    ).count()
    numero_dia = dias_anteriores + 1
    nome_dia = get_nome_dia(numero_dia)
    
    if request.method == 'POST':
        # Atualizar dados do dia
        dia.data = request.form.get('data','')
        dia.hora_ini = request.form.get('hora_ini','')
        dia.hora_fim = request.form.get('hora_fim','')
        dia.numero_ocorrencia = request.form.get('numero_ocorrencia','')
        dia.condicoes = request.form.get('condicoes','')
        dia.temp_inicial = request.form.get('temp_inicial','')
        dia.temp_final = request.form.get('temp_final','')
        if dia.tipo_busca == 'aquatica':
            dia.temp_agua = request.form.get('temp_agua','')
        else:
            dia.temp_agua = ''
        dia.guarnicao = request.form.get('guarnicao','')
        dia.recursos = request.form.get('recursos','')
        dia.historico = request.form.get('historico','')
        dia.status_vitima = request.form.get('status_vitima','')
        
        # Processar imagens (apenas se novas imagens foram enviadas)
        for campo, form_key in [
            ('img_tab_mare','img_tab_mare'),
            ('img_prev_temp','img_prev_temp'),
            ('img_traj_buscas','img_traj_buscas')
        ]:
            f = request.files.get(form_key)
            if f and f.filename:
                filename = secure_filename(f.filename)
                path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                f.save(path)
                # Comprime a imagem imediatamente após upload (pode alterar o caminho se converter formato)
                compressed_path = compress_image_on_upload(path)
                setattr(dia, campo, compressed_path)
                print(f"[upload] Atualizado dia '{campo}' em {os.path.abspath(path)}")
                log_memory_usage(f"Após atualizar dia {campo}")
        
        db.session.commit()
        flash('Dia de busca atualizado com sucesso!', 'success')
        return redirect(url_for('visualizar', id=dia.ocorrencia_id))
    
    return render_template('editar_dia.html', dia=dia, ocorrencia=ocorrencia, numero_dia=numero_dia, nome_dia=nome_dia)

@app.route('/excluir_dia/<int:id>')
def excluir_dia(id):
    dia = DiaBusca.query.get_or_404(id)
    ocorrencia_id = dia.ocorrencia_id
    
    # Excluir imagens do dia de busca
    campos_dia = ['img_tab_mare', 'img_prev_temp', 'img_traj_buscas']
    for campo in campos_dia:
        caminho_imagem = getattr(dia, campo)
        if caminho_imagem and os.path.exists(caminho_imagem):
            try:
                os.remove(caminho_imagem)
                print(f"[delete] Imagem do dia removida: {caminho_imagem}")
            except Exception as e:
                print(f"[delete] Erro ao remover imagem do dia {caminho_imagem}: {e}")
    
    db.session.delete(dia)
    db.session.commit()
    
    flash('Dia de busca excluído com sucesso!', 'success')
    return redirect(url_for('visualizar', id=ocorrencia_id))

@app.route('/excluir/<int:id>')
def excluir(id):
    ocorrencia = Ocorrencia.query.get_or_404(id)
    
    # Excluir imagens associadas
    campos_imagem = [
        'img_condic_metereologica', 'img_local', 'img_upv', 'img_satelite_upv',
        'img_raio_busca', 'img_tab_mare', 'img_prev_temp_onda'
    ]
    
    for campo in campos_imagem:
        caminho_imagem = getattr(ocorrencia, campo)
        if caminho_imagem and os.path.exists(caminho_imagem):
            try:
                os.remove(caminho_imagem)
                print(f"[delete] Imagem removida: {caminho_imagem}")
            except Exception as e:
                print(f"[delete] Erro ao remover imagem {caminho_imagem}: {e}")
    
    # Excluir imagens customizadas associadas
    imagens_customizadas = ImagemCustomizada.query.filter_by(ocorrencia_id=id).all()
    for img_custom in imagens_customizadas:
        if img_custom.caminho and os.path.exists(img_custom.caminho):
            try:
                os.remove(img_custom.caminho)
                print(f"[delete] Imagem customizada removida: {img_custom.caminho}")
            except Exception as e:
                print(f"[delete] Erro ao remover imagem customizada {img_custom.caminho}: {e}")
        db.session.delete(img_custom)
    
    # Excluir dias de busca associados
    dias_busca = DiaBusca.query.filter_by(ocorrencia_id=id).all()
    for dia in dias_busca:
        # Excluir imagens dos dias de busca
        campos_dia = ['img_tab_mare', 'img_prev_temp', 'img_traj_buscas']
        for campo in campos_dia:
            caminho_imagem = getattr(dia, campo)
            if caminho_imagem and os.path.exists(caminho_imagem):
                try:
                    os.remove(caminho_imagem)
                    print(f"[delete] Imagem do dia removida: {caminho_imagem}")
                except Exception as e:
                    print(f"[delete] Erro ao remover imagem do dia {caminho_imagem}: {e}")
        db.session.delete(dia)
    
    # Excluir relatório final associado
    relatorio_final = RelatorioFinal.query.filter_by(ocorrencia_id=id).first()
    if relatorio_final:
        # Excluir imagens do relatório final
        campos_final = ['img_corpo', 'img_local_corpo']
        for campo in campos_final:
            caminho_imagem = getattr(relatorio_final, campo)
            if caminho_imagem and os.path.exists(caminho_imagem):
                try:
                    os.remove(caminho_imagem)
                    print(f"[delete] Imagem do relatório final removida: {caminho_imagem}")
                except Exception as e:
                    print(f"[delete] Erro ao remover imagem do relatório final {caminho_imagem}: {e}")
        db.session.delete(relatorio_final)
    
    # Excluir a ocorrência
    db.session.delete(ocorrencia)
    db.session.commit()
    
    flash('Ocorrência excluída com sucesso!', 'success')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
